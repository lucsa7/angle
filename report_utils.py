# ───────────────────────────────────────────────────────────────
#  report_utils.py   ·   Generador de informes DOCX (v2-jul/25)
#  Compatible con Python 3.9+  ·  Requiere python-docx
# ───────────────────────────────────────────────────────────────
from io import BytesIO
from datetime import datetime
from dataclasses import dataclass
from typing import Tuple, Dict, List, Optional

import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ─────────── Utilidad interna: sombrear celdas ────────────────
def _shade(cell, hex_color: str = "FFFFFF"):
    """
    Aplica un fondo (hex RGB, sin #) a una celda de tabla.
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)

# ───────────── Base de datos de métricas ───────────────────────
@dataclass
class MetricInfo:
    opt: Tuple[float, float]                 # rango óptimo
    acc: Optional[Tuple[float, float]]       # rango aceptable
    type: str                                # "angle", "distance", "ratio"
    note: str                                # comentario / referencia

# --- Tabla de referencia (idéntica a la tuya, se conserva) ---
METRIC_DB: Dict[str, MetricInfo] = {
    # ---------- VISTA SAGITAL ----------
    "Hip flex": MetricInfo(opt=(45, 120),  acc=(35, 120),  type="angle",
                           note="Flexión mínima de cadera para sentadilla profunda (Clark 2019)."),
    "Knee flex": MetricInfo(opt=(90, 150), acc=(70, 150),  type="angle",
                            note="Flexión de rodilla para OHS estable (Schoenfeld 2021)."),
    "Shoulder flex": MetricInfo(opt=(150, 210), acc=(130, 210), type="angle",
                                note="Flexión de hombro para soporte de barra (McGill 2014)."),
    "|Trunk-Tibia|": MetricInfo(opt=(0, 5), acc=(0, 10), type="angle",
                                note="Alineación tronco–tibia (Glassbrook 2017)."),
    "Ankle DF": MetricInfo(opt=(15, 40), acc=(10, 40), type="angle",
                           note="Dorsiflexión reduce compensaciones (Fry 2003)."),
    # ---------- VISTA FRONTAL ----------
    "Apertura rodillas": MetricInfo(opt=(120, 9_999), acc=(80, 9_999), type="distance",
                                    note="Espacio de cadera adecuado (Escamilla 2001)."),
    "Apertura pies": MetricInfo(opt=(110, 9_999), acc=(70, 9_999), type="distance",
                                note="Base amplia mejora estabilidad (Comfort 2012)."),
    "Knee/Foot ratio": MetricInfo(opt=(0.9, 1.1), acc=(0.8, 1.2), type="ratio",
                                  note="Rodillas sobre pies (Myer 2013)."),
    "L Knee–Toe Δ (px)": MetricInfo(opt=(-12, 12), acc=(-25, 25), type="distance",
                                    note="Desplazamiento lateral ≤12 px (Padua 2012)."),
    "R Knee–Toe Δ (px)": MetricInfo(opt=(-12, 12), acc=(-25, 25), type="distance",
                                    note="Interpretación igual que lado izquierdo."),
    "Left Foot ER (°)": MetricInfo(opt=(5, 30), acc=(0, 40), type="angle",
                                   note="Rotación externa moderada (Consensus 2020)."),
    "Right Foot ER (°)": MetricInfo(opt=(5, 30), acc=(0, 40), type="angle",
                                    note="Interpretación igual que lado izquierdo."),
}

# ───────────── Interpretación de métricas ──────────────────────
def interpret_metrics(data: Dict[str, float]) -> List[Tuple[str, float, str, str]]:
    """
    Devuelve lista de tuplas:
        (métrica, valor, estado, explicación)
    Estado = 'Óptimo' · 'Aceptable' · 'Revisar' · 'Sin criterio'
    """
    results: List[Tuple[str, float, str, str]] = []

    for k, v in data.items():
        if k not in METRIC_DB:
            results.append((k, v, "Sin criterio", "Sin referencia"))
            continue

        info = METRIC_DB[k]
        estado: str

        if info.type == "angle":
            estado = (
                "Óptimo" if info.opt[0] <= v <= info.opt[1] else
                "Aceptable" if info.acc and info.acc[0] <= v <= info.acc[1] else
                "Revisar"
            )

        elif info.type == "distance" and "Apertura" in k:
            estado = (
                "Óptimo" if v >= info.opt[0] else
                "Aceptable" if v >= info.acc[0] else
                "Revisar"
            )

        elif info.type == "distance":  # absoluto
            estado = (
                "Óptimo" if info.opt[0] <= v <= info.opt[1] else
                "Aceptable" if info.acc and info.acc[0] <= v <= info.acc[1] else
                "Revisar"
            )

        elif info.type == "ratio":
            o_min, o_max = info.opt
            a_min, a_max = info.acc
            estado = (
                "Óptimo" if o_min <= v <= o_max else
                "Aceptable" if a_min <= v <= a_max else
                "Revisar"
            )

        else:
            estado = "Sin criterio"

        results.append((k, v, estado, info.note))
    return results

# ───────────── Generador de informe DOCX ───────────────────────
def build_report(img_bytes: bytes,
                 data: Dict[str, float],
                 atleta: str = "Atleta",
                 vista: str = "Sagital") -> BytesIO:
    """
    Crea y devuelve un BytesIO con un informe DOCX:
      • Portada
      • Imagen anotada (tamaño inalterado)
      • Tabla de métricas con colores
      • Recomendaciones automáticas
      • Línea de firma
    """
    doc = Document()


    # ——— Portada ———
    doc.add_heading("INFORME DE ANÁLISIS BIOMECÁNICO", level=0)
    doc.add_paragraph(f"Atleta: {atleta}")
    doc.add_paragraph(f"Vista: {vista} · Fecha: {datetime.today():%d-%m-%Y}")
    doc.add_paragraph("STA METHODOLOGIES").style = "Subtitle"

    # ——— Imagen (mismo tamaño que versión original) ———
    pic_width = Inches(1.5) if vista.lower().startswith("sag") else Inches(2.5)
    tmp = BytesIO(img_bytes); tmp.seek(0)
    doc.add_picture(tmp, width=pic_width)

    # ——— Tabla de resultados ———
    results = interpret_metrics(data)
            # ——— Glosario de interpretación ———
    doc.add_heading("Guía de interpretación", level=1)
    doc.add_paragraph((
        "Las métricas han sido clasificadas según tres niveles:\n"
        "• Óptimo: dentro del rango ideal según literatura científica.\n"
        "• Aceptable: dentro de un margen funcional pero no ideal.\n"
        "• Revisar: fuera de los rangos esperados, puede indicar limitación técnica, falta de movilidad o compensación.\n\n"
        "Cada métrica está acompañada de una breve explicación con su referencia correspondiente. "
        "Se recomienda considerar las que aparecen marcadas como 'Revisar' dentro del bloque de recomendaciones finales."
    )).style.font.size = Pt(10)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr_vals = ("Métrica", "Valor", "Estado", "Referencia")
    for i, txt in enumerate(hdr_vals):
        hdr[i].text = txt
        _shade(hdr[i], "D9D9D9")                     # gris header
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True

    color_map = {"Óptimo": "C6EFCE",    # verde claro
                 "Aceptable": "FFEB9C", # amarillo
                 "Revisar": "F8CBAD"}   # rojo claro

    for met, val, est, note in results:
        cells = table.add_row().cells
        cells[0].text = met
        cells[1].text = f"{val:.1f}" if isinstance(val, (int, float, np.floating)) else str(val)
        cells[2].text = est
        cells[3].text = note

        # Sombreado según estado
        _shade(cells[2], color_map.get(est, "FFFFFF"))

        # Tamaño de fuente uniforme
        for c in cells:
            for run in c.paragraphs[0].runs:
                run.font.size = Pt(10)

    # ——— Recomendaciones ———
    revisar = [r for r in results if r[2] == "Revisar"]
    doc.add_heading("Recomendaciones", level=1)

    if revisar:
        for met, val, est, note in revisar:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{met}: ").bold = True
            p.add_run(note)
    else:
        doc.add_paragraph("No se detectaron métricas fuera de los rangos óptimo o aceptable.")

    # ——— Firma ———
    doc.add_paragraph("\n\n_____________________________\n")

    # ——— Salida en memoria ———
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

